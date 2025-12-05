"""
Script pour générer la documentation de la base de données en format Word (.docx)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Définir la couleur de fond d'une cellule"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_database_documentation():
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('SYGLA-H2O - Structure de la Base de Données', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    subtitle = doc.add_paragraph('Documentation complète des tables PostgreSQL')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Généré le: 3 décembre 2025')
    doc.add_paragraph()
    
    # ========================================
    # TABLE DES MATIÈRES
    # ========================================
    doc.add_heading('Table des Matières', level=1)
    toc = [
        '1. Authentification & Sécurité',
        '   - authentication_user',
        '   - security_settings',
        '   - user_sessions',
        '2. Gestion des Clients',
        '   - clients_client',
        '3. Gestion des Produits & Stock',
        '   - products_produit',
        '   - products_mouvementstock',
        '4. Gestion des Commandes',
        '   - orders_commande',
        '   - orders_itemcommande',
        '   - paiements_commande',
        '5. Gestion des Ventes',
        '   - ventes',
        '   - lignes_vente',
        '   - paiements',
        '6. Notifications',
        '   - notifications',
        '   - notification_preferences',
        '7. Logs Système',
        '   - logs_systemlog',
    ]
    for item in toc:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========================================
    # TABLES
    # ========================================
    
    tables_data = {
        'AUTHENTICATION_USER': {
            'description': 'Table principale des utilisateurs du système. Stocke les informations d\'authentification et les profils.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique auto-incrémenté'),
                ('username', 'VARCHAR(150)', 'UNIQUE, NOT NULL', 'Nom d\'utilisateur unique'),
                ('password', 'VARCHAR(128)', 'NOT NULL', 'Mot de passe hashé'),
                ('email', 'VARCHAR(254)', 'UNIQUE, NOT NULL', 'Adresse email unique'),
                ('first_name', 'VARCHAR(150)', '', 'Prénom'),
                ('last_name', 'VARCHAR(150)', '', 'Nom de famille'),
                ('role', 'VARCHAR(20)', 'NOT NULL', 'Rôle: admin, vendeur, stock, livreur'),
                ('telephone', 'VARCHAR(17)', '', 'Numéro de téléphone'),
                ('adresse', 'TEXT', '', 'Adresse physique'),
                ('photo', 'VARCHAR(100)', 'NULL', 'Chemin vers la photo de profil'),
                ('is_active', 'BOOLEAN', 'DEFAULT TRUE', 'Compte actif ou désactivé'),
                ('is_staff', 'BOOLEAN', 'DEFAULT FALSE', 'Accès à l\'admin Django'),
                ('is_superuser', 'BOOLEAN', 'DEFAULT FALSE', 'Super utilisateur'),
                ('must_change_password', 'BOOLEAN', 'DEFAULT FALSE', 'Forcer changement mot de passe'),
                ('last_login', 'DATETIME', 'NULL', 'Dernière connexion'),
                ('last_activity', 'DATETIME', 'NULL', 'Dernière activité'),
                ('date_joined', 'DATETIME', 'NOT NULL', 'Date d\'inscription'),
                ('date_creation', 'DATETIME', 'NOT NULL', 'Date de création'),
                ('date_modification', 'DATETIME', 'NOT NULL', 'Dernière modification'),
            ]
        },
        'CLIENTS_CLIENT': {
            'description': 'Table des clients (entreprises commerciales). Gère les informations clients et leur crédit.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('nom_commercial', 'VARCHAR(200)', 'NULL', 'Nom commercial de l\'entreprise'),
                ('raison_sociale', 'VARCHAR(200)', 'NULL', 'Raison sociale légale'),
                ('telephone', 'VARCHAR(17)', 'NOT NULL', 'Téléphone principal'),
                ('email', 'VARCHAR(254)', '', 'Email de contact'),
                ('adresse', 'TEXT', '', 'Adresse complète'),
                ('contact', 'VARCHAR(200)', '', 'Nom du contact principal'),
                ('credit_limite', 'DECIMAL(10,2)', 'DEFAULT 0', 'Limite de crédit accordée'),
                ('credit_utilise', 'DECIMAL(10,2)', 'DEFAULT 0', 'Crédit actuellement utilisé'),
                ('notes', 'TEXT', '', 'Notes et commentaires'),
                ('is_active', 'BOOLEAN', 'DEFAULT TRUE', 'Client actif'),
                ('date_creation', 'DATETIME', 'NOT NULL', 'Date de création'),
                ('date_modification', 'DATETIME', 'NOT NULL', 'Dernière modification'),
            ]
        },
        'PRODUCTS_PRODUIT': {
            'description': 'Catalogue des produits (eau et glace). Gère les prix et le stock.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('code_produit', 'VARCHAR(20)', 'UNIQUE, NOT NULL', 'Code unique du produit'),
                ('nom', 'VARCHAR(200)', 'NOT NULL', 'Nom du produit'),
                ('description', 'TEXT', '', 'Description détaillée'),
                ('type_produit', 'VARCHAR(50)', 'NOT NULL', 'Type: eau, glace'),
                ('unite_mesure', 'VARCHAR(50)', 'NOT NULL', 'Unité: gallon, livre, etc.'),
                ('prix_unitaire', 'DECIMAL(10,2)', 'NOT NULL', 'Prix de vente unitaire'),
                ('stock_actuel', 'INTEGER', 'DEFAULT 0', 'Quantité en stock'),
                ('stock_minimal', 'INTEGER', 'DEFAULT 0', 'Seuil d\'alerte stock bas'),
                ('stock_initial', 'INTEGER', 'DEFAULT 0', 'Stock initial'),
                ('is_active', 'BOOLEAN', 'DEFAULT TRUE', 'Produit actif'),
                ('date_creation', 'DATETIME', 'NOT NULL', 'Date de création'),
                ('date_modification', 'DATETIME', 'NOT NULL', 'Dernière modification'),
            ]
        },
        'PRODUCTS_MOUVEMENTSTOCK': {
            'description': 'Historique de tous les mouvements de stock (entrées, sorties, ajustements).',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('produit_id', 'INTEGER', 'FK → products_produit', 'Référence au produit'),
                ('utilisateur_id', 'INTEGER', 'FK → authentication_user', 'Utilisateur responsable'),
                ('type_mouvement', 'VARCHAR(20)', 'NOT NULL', 'Type: entree, sortie, ajustement'),
                ('quantite', 'INTEGER', 'NOT NULL', 'Quantité du mouvement'),
                ('stock_avant', 'INTEGER', 'NOT NULL', 'Stock avant le mouvement'),
                ('stock_apres', 'INTEGER', 'NOT NULL', 'Stock après le mouvement'),
                ('motif', 'VARCHAR(200)', '', 'Raison du mouvement'),
                ('numero_document', 'VARCHAR(100)', '', 'N° de document associé'),
                ('date_creation', 'DATETIME', 'NOT NULL', 'Date du mouvement'),
            ]
        },
        'ORDERS_COMMANDE': {
            'description': 'Commandes clients avec suivi complet du workflow (création → validation → livraison).',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('numero_commande', 'VARCHAR(20)', 'UNIQUE', 'Numéro de commande généré'),
                ('client_id', 'INTEGER', 'FK → clients_client', 'Client de la commande'),
                ('vendeur_id', 'INTEGER', 'FK → authentication_user', 'Vendeur créateur'),
                ('statut', 'VARCHAR(20)', 'NOT NULL', 'Statut: en_attente, validee, en_preparation, en_livraison, livree, annulee'),
                ('montant_produits', 'DECIMAL(10,2)', 'NOT NULL', 'Sous-total des produits'),
                ('frais_livraison', 'DECIMAL(10,2)', 'DEFAULT 0', 'Frais de livraison'),
                ('montant_total', 'DECIMAL(10,2)', 'NOT NULL', 'Montant total TTC'),
                ('montant_paye', 'DECIMAL(10,2)', 'DEFAULT 0', 'Montant déjà payé'),
                ('montant_restant', 'DECIMAL(10,2)', 'DEFAULT 0', 'Reste à payer'),
                ('statut_paiement', 'VARCHAR(20)', 'NOT NULL', 'Statut: impaye, paye_partiel, paye'),
                ('type_livraison', 'VARCHAR(20)', 'NOT NULL', 'Type: livraison, pickup'),
                ('livreur', 'VARCHAR(100)', 'NULL', 'Nom du livreur assigné'),
                ('adresse_livraison', 'TEXT', '', 'Adresse de livraison'),
                ('instructions_livraison', 'TEXT', '', 'Instructions spéciales'),
                ('notes', 'TEXT', '', 'Notes internes'),
                ('date_creation', 'DATETIME', 'NOT NULL', 'Date de création'),
                ('date_validation', 'DATETIME', 'NULL', 'Date de validation'),
                ('date_livraison_prevue', 'DATETIME', 'NULL', 'Date prévue de livraison'),
                ('date_livraison_effective', 'DATETIME', 'NULL', 'Date réelle de livraison'),
                ('date_echeance', 'DATE', 'NULL', 'Date d\'échéance paiement'),
                ('convertie_en_vente', 'BOOLEAN', 'DEFAULT FALSE', 'Commande convertie en vente'),
                ('vente_associee_id', 'INTEGER', 'FK → ventes', 'Vente liée'),
            ]
        },
        'ORDERS_ITEMCOMMANDE': {
            'description': 'Lignes de détail des commandes (produits commandés).',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('commande_id', 'INTEGER', 'FK → orders_commande', 'Référence à la commande'),
                ('produit_id', 'INTEGER', 'FK → products_produit', 'Produit commandé'),
                ('quantite', 'INTEGER', 'NOT NULL', 'Quantité commandée'),
                ('prix_unitaire', 'DECIMAL(10,2)', 'NOT NULL', 'Prix unitaire au moment de la commande'),
                ('sous_total', 'DECIMAL(10,2)', 'NOT NULL', 'Sous-total de la ligne'),
            ]
        },
        'PAIEMENTS_COMMANDE': {
            'description': 'Paiements effectués sur les commandes.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('commande_id', 'INTEGER', 'FK → orders_commande', 'Commande payée'),
                ('recu_par_id', 'INTEGER', 'FK → authentication_user', 'Utilisateur qui a reçu le paiement'),
                ('montant', 'DECIMAL(10,2)', 'NOT NULL', 'Montant du paiement'),
                ('methode', 'VARCHAR(20)', 'NOT NULL', 'Méthode: especes, carte, virement, cheque, mobile'),
                ('reference', 'VARCHAR(100)', '', 'Référence du paiement'),
                ('notes', 'TEXT', '', 'Notes sur le paiement'),
                ('date_paiement', 'DATETIME', 'NOT NULL', 'Date et heure du paiement'),
            ]
        },
        'VENTES': {
            'description': 'Ventes finalisées. Créées directement ou à partir de commandes validées.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('numero_vente', 'VARCHAR(50)', 'UNIQUE, NOT NULL', 'Numéro de vente unique'),
                ('client_id', 'INTEGER', 'FK → clients_client', 'Client de la vente'),
                ('vendeur_id', 'INTEGER', 'FK → authentication_user', 'Vendeur'),
                ('montant_total', 'DECIMAL(10,2)', 'NOT NULL', 'Montant total TTC'),
                ('montant_paye', 'DECIMAL(10,2)', 'DEFAULT 0', 'Montant payé'),
                ('montant_restant', 'DECIMAL(10,2)', 'DEFAULT 0', 'Reste à payer'),
                ('statut_paiement', 'VARCHAR(20)', 'NOT NULL', 'Statut: impaye, paye_partiel, paye'),
                ('methode_paiement', 'VARCHAR(20)', '', 'Méthode principale'),
                ('remise_pourcentage', 'DECIMAL(5,2)', 'DEFAULT 0', 'Pourcentage de remise'),
                ('remise_montant', 'DECIMAL(10,2)', 'DEFAULT 0', 'Montant de remise'),
                ('type_livraison', 'VARCHAR(50)', '', 'Type de livraison'),
                ('frais_livraison', 'DECIMAL(10,2)', 'DEFAULT 0', 'Frais de livraison'),
                ('frais_supplementaires', 'DECIMAL(10,2)', 'DEFAULT 0', 'Autres frais'),
                ('raison_frais', 'VARCHAR(200)', '', 'Justification des frais'),
                ('notes', 'TEXT', '', 'Notes'),
                ('date_vente', 'DATETIME', 'NOT NULL', 'Date de la vente'),
                ('date_echeance', 'DATE', 'NULL', 'Échéance paiement'),
                ('date_livraison_prevue', 'DATE', 'NULL', 'Date de livraison prévue'),
                ('created_at', 'DATETIME', 'NOT NULL', 'Création'),
                ('updated_at', 'DATETIME', 'NOT NULL', 'Modification'),
            ]
        },
        'LIGNES_VENTE': {
            'description': 'Détail des produits vendus par vente.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('vente_id', 'INTEGER', 'FK → ventes', 'Référence à la vente'),
                ('produit_id', 'INTEGER', 'FK → products_produit', 'Produit vendu'),
                ('quantite', 'DECIMAL(10,2)', 'NOT NULL', 'Quantité vendue'),
                ('prix_unitaire', 'DECIMAL(10,2)', 'NOT NULL', 'Prix unitaire'),
                ('montant', 'DECIMAL(10,2)', 'NOT NULL', 'Montant de la ligne'),
            ]
        },
        'PAIEMENTS': {
            'description': 'Paiements effectués sur les ventes.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('vente_id', 'INTEGER', 'FK → ventes', 'Vente payée'),
                ('recu_par_id', 'INTEGER', 'FK → authentication_user', 'Reçu par'),
                ('montant', 'DECIMAL(10,2)', 'NOT NULL', 'Montant'),
                ('methode', 'VARCHAR(20)', 'NOT NULL', 'Méthode de paiement'),
                ('reference', 'VARCHAR(100)', '', 'Référence'),
                ('notes', 'TEXT', '', 'Notes'),
                ('date_paiement', 'DATETIME', 'NOT NULL', 'Date du paiement'),
            ]
        },
        'NOTIFICATIONS': {
            'description': 'Notifications système envoyées aux utilisateurs.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('user_id', 'INTEGER', 'FK → authentication_user', 'Destinataire'),
                ('type', 'VARCHAR(50)', 'NOT NULL', 'Type de notification'),
                ('title', 'VARCHAR(200)', 'NOT NULL', 'Titre'),
                ('message', 'TEXT', 'NOT NULL', 'Contenu'),
                ('related_order_id', 'INTEGER', 'NULL', 'Commande liée'),
                ('related_product_id', 'INTEGER', 'NULL', 'Produit lié'),
                ('related_client_id', 'INTEGER', 'NULL', 'Client lié'),
                ('related_sale_id', 'INTEGER', 'NULL', 'Vente liée'),
                ('is_read', 'BOOLEAN', 'DEFAULT FALSE', 'Lu ou non'),
                ('read_at', 'DATETIME', 'NULL', 'Date de lecture'),
                ('created_at', 'DATETIME', 'NOT NULL', 'Date de création'),
            ]
        },
        'NOTIFICATION_PREFERENCES': {
            'description': 'Préférences de notification par utilisateur.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('user_id', 'INTEGER', 'FK UNIQUE → authentication_user', 'Utilisateur'),
                ('notify_client_created', 'BOOLEAN', 'DEFAULT TRUE', 'Notification nouveau client'),
                ('notify_order_created', 'BOOLEAN', 'DEFAULT TRUE', 'Notification nouvelle commande'),
                ('notify_order_validated', 'BOOLEAN', 'DEFAULT TRUE', 'Notification commande validée'),
                ('notify_delivery_assigned', 'BOOLEAN', 'DEFAULT TRUE', 'Notification livraison assignée'),
                ('notify_delivery_completed', 'BOOLEAN', 'DEFAULT TRUE', 'Notification livraison terminée'),
                ('notify_stock_low', 'BOOLEAN', 'DEFAULT TRUE', 'Alerte stock bas'),
                ('notify_stock_updated', 'BOOLEAN', 'DEFAULT TRUE', 'Mise à jour stock'),
                ('notify_system_errors', 'BOOLEAN', 'DEFAULT TRUE', 'Erreurs système'),
                ('notify_security_alerts', 'BOOLEAN', 'DEFAULT TRUE', 'Alertes sécurité'),
                ('notify_daily_report', 'BOOLEAN', 'DEFAULT FALSE', 'Rapport quotidien'),
                ('notify_weekly_report', 'BOOLEAN', 'DEFAULT FALSE', 'Rapport hebdomadaire'),
                ('notify_monthly_report', 'BOOLEAN', 'DEFAULT FALSE', 'Rapport mensuel'),
                ('enable_email_notifications', 'BOOLEAN', 'DEFAULT TRUE', 'Notifications email'),
                ('enable_browser_notifications', 'BOOLEAN', 'DEFAULT TRUE', 'Notifications navigateur'),
                ('created_at', 'DATETIME', 'NOT NULL', 'Création'),
                ('updated_at', 'DATETIME', 'NOT NULL', 'Modification'),
            ]
        },
        'SECURITY_SETTINGS': {
            'description': 'Paramètres de sécurité globaux du système.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('session_timeout', 'INTEGER', 'DEFAULT 30', 'Timeout session (minutes)'),
                ('max_login_attempts', 'INTEGER', 'DEFAULT 5', 'Tentatives max avant blocage'),
                ('lockout_duration', 'INTEGER', 'DEFAULT 30', 'Durée blocage (minutes)'),
                ('password_min_length', 'INTEGER', 'DEFAULT 8', 'Longueur min mot de passe'),
                ('require_uppercase', 'BOOLEAN', 'DEFAULT TRUE', 'Exiger majuscule'),
                ('require_lowercase', 'BOOLEAN', 'DEFAULT TRUE', 'Exiger minuscule'),
                ('require_numbers', 'BOOLEAN', 'DEFAULT TRUE', 'Exiger chiffres'),
                ('require_special_chars', 'BOOLEAN', 'DEFAULT FALSE', 'Exiger caractères spéciaux'),
                ('jwt_access_token_lifetime', 'INTEGER', 'DEFAULT 60', 'Durée token access (minutes)'),
                ('jwt_refresh_token_lifetime', 'INTEGER', 'DEFAULT 1440', 'Durée token refresh (minutes)'),
                ('enable_two_factor', 'BOOLEAN', 'DEFAULT FALSE', 'Activer 2FA'),
                ('force_password_change', 'INTEGER', 'DEFAULT 90', 'Forcer changement (jours)'),
                ('updated_by_id', 'INTEGER', 'FK → authentication_user', 'Modifié par'),
                ('updated_at', 'DATETIME', 'NOT NULL', 'Dernière modification'),
            ]
        },
        'USER_SESSIONS': {
            'description': 'Sessions utilisateur actives pour le suivi des connexions.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('user_id', 'INTEGER', 'FK → authentication_user', 'Utilisateur'),
                ('token', 'VARCHAR(500)', 'UNIQUE, NOT NULL', 'Token de session'),
                ('ip_address', 'VARCHAR(39)', 'NULL', 'Adresse IP'),
                ('user_agent', 'TEXT', '', 'User agent du navigateur'),
                ('device_info', 'VARCHAR(200)', '', 'Info sur l\'appareil'),
                ('login_time', 'DATETIME', 'NOT NULL', 'Heure de connexion'),
                ('last_activity', 'DATETIME', 'NOT NULL', 'Dernière activité'),
                ('is_active', 'BOOLEAN', 'DEFAULT TRUE', 'Session active'),
                ('logout_time', 'DATETIME', 'NULL', 'Heure de déconnexion'),
            ]
        },
        'LOGS_SYSTEMLOG': {
            'description': 'Logs système pour le suivi des actions et erreurs.',
            'fields': [
                ('id', 'INTEGER', 'PK', 'Identifiant unique'),
                ('user_id', 'INTEGER', 'FK → authentication_user', 'Utilisateur concerné'),
                ('type', 'VARCHAR(20)', 'NOT NULL', 'Type: info, warning, error, success'),
                ('message', 'VARCHAR(500)', 'NOT NULL', 'Message du log'),
                ('details', 'TEXT', '', 'Détails complets'),
                ('module', 'VARCHAR(50)', '', 'Module concerné'),
                ('ip_address', 'VARCHAR(39)', 'NULL', 'Adresse IP'),
                ('user_agent', 'TEXT', '', 'User agent'),
                ('request_method', 'VARCHAR(10)', '', 'Méthode HTTP'),
                ('endpoint', 'VARCHAR(255)', '', 'Endpoint appelé'),
                ('status_code', 'INTEGER', 'NULL', 'Code HTTP'),
                ('response_time', 'VARCHAR(50)', '', 'Temps de réponse'),
                ('metadata', 'JSON', '', 'Données additionnelles'),
                ('timestamp', 'DATETIME', 'NOT NULL', 'Date et heure'),
            ]
        },
    }
    
    # Générer les sections pour chaque table
    for table_name, table_info in tables_data.items():
        doc.add_heading(table_name, level=1)
        
        # Description
        doc.add_paragraph(table_info['description'])
        doc.add_paragraph()
        
        # Créer le tableau
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # En-têtes
        header_cells = table.rows[0].cells
        headers = ['Champ', 'Type', 'Contraintes', 'Description']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(header_cells[i], '4472C4')
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Données
        for idx, field in enumerate(table_info['fields']):
            row = table.add_row()
            for i, value in enumerate(field):
                row.cells[i].text = str(value)
                # Alterner les couleurs des lignes
                if idx % 2 == 0:
                    set_cell_shading(row.cells[i], 'D9E2F3')
        
        doc.add_paragraph()
        doc.add_page_break()
    
    # ========================================
    # RELATIONS
    # ========================================
    doc.add_heading('Diagramme des Relations', level=1)
    
    relations = [
        ('orders_commande.client_id', '→', 'clients_client.id'),
        ('orders_commande.vendeur_id', '→', 'authentication_user.id'),
        ('orders_commande.vente_associee_id', '→', 'ventes.id'),
        ('orders_itemcommande.commande_id', '→', 'orders_commande.id'),
        ('orders_itemcommande.produit_id', '→', 'products_produit.id'),
        ('paiements_commande.commande_id', '→', 'orders_commande.id'),
        ('paiements_commande.recu_par_id', '→', 'authentication_user.id'),
        ('ventes.client_id', '→', 'clients_client.id'),
        ('ventes.vendeur_id', '→', 'authentication_user.id'),
        ('lignes_vente.vente_id', '→', 'ventes.id'),
        ('lignes_vente.produit_id', '→', 'products_produit.id'),
        ('paiements.vente_id', '→', 'ventes.id'),
        ('paiements.recu_par_id', '→', 'authentication_user.id'),
        ('products_mouvementstock.produit_id', '→', 'products_produit.id'),
        ('products_mouvementstock.utilisateur_id', '→', 'authentication_user.id'),
        ('notifications.user_id', '→', 'authentication_user.id'),
        ('notification_preferences.user_id', '→', 'authentication_user.id'),
        ('user_sessions.user_id', '→', 'authentication_user.id'),
        ('logs_systemlog.user_id', '→', 'authentication_user.id'),
    ]
    
    rel_table = doc.add_table(rows=1, cols=3)
    rel_table.style = 'Table Grid'
    
    header_cells = rel_table.rows[0].cells
    header_cells[0].text = 'Clé Étrangère'
    header_cells[1].text = ''
    header_cells[2].text = 'Référence'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    for rel in relations:
        row = rel_table.add_row()
        for i, value in enumerate(rel):
            row.cells[i].text = value
    
    # Sauvegarder
    output_path = 'SYGLA_H2O_Structure_Base_Donnees.docx'
    doc.save(output_path)
    print(f'✅ Document créé avec succès: {output_path}')
    return output_path

if __name__ == '__main__':
    create_database_documentation()
